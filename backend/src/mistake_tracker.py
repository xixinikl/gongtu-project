"""错题追踪模块：存储、分析、导出"""
import json
import os
import uuid
import logging
from datetime import datetime, timezone, timedelta
from threading import Lock
from typing import Optional

from openai import OpenAI

from src.models import MistakeRecord, MistakeAnalysis

logger = logging.getLogger("mistake_tracker")

DATA_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data")
MISTAKES_FILE = os.path.join(DATA_DIR, "mistakes.json")
EXPORT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "exports")

_lock = Lock()

# ── LLM 配置 ──
API_KEY = os.getenv("LLM_API_KEY", "")
BASE_URL = os.getenv("LLM_BASE_URL", "https://api.deepseek.com")
MODEL = os.getenv("LLM_MODEL", "deepseek-chat")
TIMEOUT = int(os.getenv("LLM_TIMEOUT", "30"))


def _load_mistakes() -> list[dict]:
    if not os.path.exists(MISTAKES_FILE):
        return []
    with open(MISTAKES_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


def _save_mistakes(data: list[dict]):
    os.makedirs(DATA_DIR, exist_ok=True)
    with open(MISTAKES_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def add_mistake(question_id: str, question_title: str, question_type: str,
                student_answer: str, ai_reply: str) -> MistakeRecord:
    """记录一条错题/批改记录。"""
    record = {
        "id": uuid.uuid4().hex[:8],
        "questionId": question_id,
        "questionTitle": question_title,
        "questionType": question_type,
        "studentAnswer": student_answer,
        "aiReply": ai_reply,
        "createdAt": datetime.now(timezone(timedelta(hours=8))).isoformat(),
    }

    with _lock:
        data = _load_mistakes()
        data.append(record)
        _save_mistakes(data)

    logger.info(f"Mistake recorded: {record['id']}, q={question_id}")
    return MistakeRecord(**record)


def get_all_mistakes() -> list[MistakeRecord]:
    """获取所有错题记录（按时间倒序）。"""
    data = _load_mistakes()
    data.sort(key=lambda x: x["createdAt"], reverse=True)
    return [MistakeRecord(**r) for r in data]


def delete_mistake(mistake_id: str) -> bool:
    """删除一条错题记录。"""
    with _lock:
        data = _load_mistakes()
        new_data = [r for r in data if r["id"] != mistake_id]
        if len(new_data) == len(data):
            return False
        _save_mistakes(new_data)
    return True


def clear_all_mistakes():
    """清空所有错题记录。"""
    with _lock:
        _save_mistakes([])


def analyze_with_ai() -> MistakeAnalysis:
    """调用 AI 分析所有错题记录，总结薄弱点。"""
    records = get_all_mistakes()
    if not records:
        return MistakeAnalysis(
            summary="暂无批改记录，无法分析。",
            weakDimensions=[],
            recommendations=["先做几道题，再来看你的薄弱点分析吧！"],
            recordsReviewed=0,
        )

    # 构建分析prompt
    records_text = ""
    for i, r in enumerate(records):
        records_text += f"\n### 记录{i+1}：{r.questionTitle}（{r.questionType}）\n"
        records_text += f"学生作答：{r.studentAnswer}\n"
        records_text += f"老师批改：{r.aiReply[:500]}\n"

    analysis_prompt = f"""请分析以下学生的申论作答批改记录，总结薄弱点和改进建议。

{records_text}

请从以下角度分析：
1. 总体薄弱点：学生在哪些方面存在问题？是否有共性问题？
2. 薄弱维度：在内容完整性、逻辑结构、语言表达、对策可行性、格式规范中，哪些维度最弱？
3. 改进建议：3-5条具体的、可操作的建议

请直接回复，用中文，不要用markdown标题。"""

    if not API_KEY or API_KEY == "your-api-key-here":
        raise RuntimeError("LLM API key not configured.")

    client = OpenAI(api_key=API_KEY, base_url=BASE_URL, timeout=TIMEOUT)

    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {"role": "system", "content": "你是飞扬老师，请分析学生的错题记录。"},
            {"role": "user", "content": analysis_prompt},
        ],
        temperature=0.5,
        max_tokens=1500,
    )

    ai_reply = response.choices[0].message.content
    if not isinstance(ai_reply, str) or not ai_reply.strip():
        raise RuntimeError("LLM returned an empty response")

    # 从AI回复中提取弱维度（简单关键词匹配）
    weak_keywords = []
    dims = ["内容完整性", "逻辑结构", "语言表达", "对策可行性", "格式规范"]
    for dim in dims:
        if dim in ai_reply and ("薄弱" in ai_reply or "较差" in ai_reply or "一般" in ai_reply or "不足" in ai_reply):
            weak_keywords.append(dim)

    # 提取建议（按序号分行）
    recommendations = []
    for line in ai_reply.split("\n"):
        line = line.strip()
        if line and (line[0].isdigit() or line.startswith("- ")):
            recommendations.append(line.lstrip("0123456789.、- ").strip())

    if not recommendations:
        recommendations = ["请根据分析结果制定针对性练习计划。"]

    return MistakeAnalysis(
        summary=ai_reply,
        weakDimensions=weak_keywords or ["待进一步确认"],
        recommendations=recommendations[:5],
        recordsReviewed=len(records),
    )


def export_to_word(analysis: Optional[MistakeAnalysis] = None) -> str:
    """导出错题分析和记录为Word文档，返回文件路径。"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    os.makedirs(EXPORT_DIR, exist_ok=True)

    doc = Document()

    # 标题
    title = doc.add_heading("飞扬老师 · 申论错题追踪报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"生成时间：{datetime.now(timezone(timedelta(hours=8))).strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    # 分析部分
    if analysis and analysis.recordsReviewed > 0:
        doc.add_heading("一、薄弱点分析", level=1)
        doc.add_paragraph(analysis.summary)

        doc.add_heading("薄弱维度", level=2)
        for dim in analysis.weakDimensions:
            p = doc.add_paragraph(f"🔴 {dim}", style="List Bullet")
            for run in p.runs:
                run.font.color.rgb = RGBColor(200, 50, 50)

        doc.add_heading("改进建议", level=2)
        for rec in analysis.recommendations:
            doc.add_paragraph(rec, style="List Bullet")

        doc.add_paragraph("")

    # 批改记录列表
    doc.add_heading("二、批改记录", level=1)
    records = get_all_mistakes()
    if records:
        for i, r in enumerate(records):
            doc.add_heading(f"记录 {i+1}：{r.questionTitle}", level=2)
            doc.add_paragraph(f"题型：{r.questionType}  |  时间：{r.createdAt[:16]}")

            doc.add_heading("学生作答", level=3)
            para = doc.add_paragraph(r.studentAnswer)
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.name = "Courier New"

            doc.add_heading("老师批改", level=3)
            doc.add_paragraph(r.aiReply)

            doc.add_paragraph("─" * 40)
    else:
        doc.add_paragraph("暂无记录。")

    # 保存
    filename = f"错题追踪_{datetime.now(timezone(timedelta(hours=8))).strftime('%Y%m%d_%H%M')}.docx"
    filepath = os.path.join(EXPORT_DIR, filename)
    doc.save(filepath)
    logger.info(f"Word exported to {filepath}")
    return filepath
